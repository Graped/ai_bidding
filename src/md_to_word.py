from pathlib import Path
import markdown
from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
import re
import subprocess
import tempfile
import os

def convert_mermaid_to_image(mermaid_code):
    """将 Mermaid 代码转换为图片"""
    # 创建临时文件
    with tempfile.NamedTemporaryFile(suffix='.mmd', delete=False, mode='w', encoding='utf-8') as f:
        # 添加主题和样式设置
        mermaid_config = """
%%{init: {'theme': 'default', 'themeVariables': { 'fontSize': '16px', 'fontFamily': '宋体' }}}%%
"""
        f.write(mermaid_config + mermaid_code)
        mmd_file = f.name
    
    # 创建输出图片文件
    png_file = mmd_file.replace('.mmd', '.png')
    
    try:
        # 使用 mmdc 命令转换，设置统一的图片大小和背景
        subprocess.run([
            'mmdc',
            '-i', mmd_file,
            '-o', png_file,
            '-w', '800',  # 设置宽度
            '-H', '600',  # 设置高度
            '-b', 'transparent',  # 设置透明背景
            '-s', '3',  # 设置缩放比例
            '-c', 'config.json'  # 使用配置文件
        ], check=True)
        return png_file
    except subprocess.CalledProcessError as e:
        print(f"转换流程图失败: {e}")
        return None
    finally:
        # 清理临时文件
        if os.path.exists(mmd_file):
            os.unlink(mmd_file)

def create_mermaid_config():
    """创建 Mermaid 配置文件"""
    config = {
        "theme": "default",
        "themeVariables": {
            "fontSize": "16px",
            "fontFamily": "宋体",
            "primaryColor": "#1f77b4",
            "primaryTextColor": "#000000",
            "primaryBorderColor": "#1f77b4",
            "lineColor": "#1f77b4",
            "secondaryColor": "#ff7f0e",
            "tertiaryColor": "#2ca02c"
        },
        "flowchart": {
            "curve": "basis",
            "padding": 15,
            "nodeSpacing": 50,
            "rankSpacing": 50
        }
    }
    
    with open('config.json', 'w', encoding='utf-8') as f:
        import json
        json.dump(config, f, indent=2)

def process_mermaid(doc, mermaid_code):
    """处理 Mermaid 流程图"""
    # 转换 Mermaid 代码为图片
    png_file = convert_mermaid_to_image(mermaid_code)
    if png_file and os.path.exists(png_file):
        try:
            # 添加图片到文档
            doc.add_picture(png_file, width=Inches(6))  # 设置统一宽度为6英寸
            
            # 设置图片居中
            last_paragraph = doc.paragraphs[-1]
            last_paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            
            # 添加图片说明（可选）
            caption = doc.add_paragraph()
            caption.alignment = WD_ALIGN_PARAGRAPH.CENTER
            caption_run = caption.add_run("图 X-X 流程图")
            caption_run.font.name = '宋体'
            caption_run.font.size = Pt(10.5)
        finally:
            # 清理临时图片文件
            os.unlink(png_file)

def set_document_styles(doc):
    """设置文档样式"""
    # 设置默认字体
    styles = doc.styles
    style = styles['Normal']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    # 设置标题样式
    for i in range(1, 5):
        style = styles[f'Heading {i}']
        style.font.name = '黑体'
        style.font.size = Pt(16 - i)  # 标题字号递减
        if i == 1:
            style.font.bold = True
    
    # 设置列表样式
    style = styles['List Bullet']
    style.font.name = '宋体'
    style.font.size = Pt(12)
    
    style = styles['List Number']
    style.font.name = '宋体'
    style.font.size = Pt(12)

def set_document_format(doc):
    """设置文档格式"""
    # 设置页面边距
    sections = doc.sections
    for section in sections:
        section.top_margin = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin = Cm(3.18)
        section.right_margin = Cm(3.18)
        
        # 添加页眉
        header = section.header
        header_para = header.paragraphs[0]
        header_para.text = "智慧工程项目管理平台智能化软件、硬件系统投标文件"
        header_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 添加页脚
        footer = section.footer
        footer_para = footer.paragraphs[0]
        footer_para.text = "第 {PAGE} 页，共 {NUMPAGES} 页"
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER

def process_table(md_table, doc):
    """处理 Markdown 表格"""
    lines = md_table.strip().split('\n')
    if len(lines) < 3:  # 至少需要表头、分隔行和一行数据
        return
    
    # 计算列数
    header_cells = lines[0].strip('|').split('|')
    col_count = len(header_cells)
    
    # 创建表格
    table = doc.add_table(rows=1, cols=col_count)
    table.style = 'Table Grid'
    
    # 添加表头
    header_row = table.rows[0]
    for i, cell in enumerate(header_cells):
        header_row.cells[i].text = cell.strip()
        # 设置表头格式
        for paragraph in header_row.cells[i].paragraphs:
            paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in paragraph.runs:
                run.bold = True
                run.font.name = '黑体'
    
    # 添加数据行
    for line in lines[2:]:  # 跳过表头和分隔行
        cells = line.strip('|').split('|')
        if len(cells) == col_count:
            row = table.add_row()
            for i, cell in enumerate(cells):
                row.cells[i].text = cell.strip()
                # 设置单元格格式
                for paragraph in row.cells[i].paragraphs:
                    paragraph.alignment = WD_ALIGN_PARAGRAPH.CENTER
                    for run in paragraph.runs:
                        run.font.name = '宋体'

def convert_md_to_word():
    # 创建 Mermaid 配置文件
    create_mermaid_config()
    
    # 读取 Markdown 文件
    md_file = Path("data/output/智慧工程项目管理平台智能化软件、硬件系统招标文件/完整投标文件.md")
    with open(md_file, 'r', encoding='utf-8') as f:
        md_content = f.read()
    
    # 创建 Word 文档
    doc = Document()
    
    # 设置文档样式和格式
    set_document_styles(doc)
    set_document_format(doc)
    
    # 处理内容
    lines = md_content.split('\n')
    i = 0
    while i < len(lines):
        line = lines[i]
        
        # 处理 Mermaid 流程图
        if line.strip().startswith('```mermaid'):
            mermaid_lines = []
            i += 1
            while i < len(lines) and not lines[i].strip().startswith('```'):
                mermaid_lines.append(lines[i])
                i += 1
            if mermaid_lines:
                process_mermaid(doc, '\n'.join(mermaid_lines))
            i += 1
            continue
        
        # 处理表格
        if line.strip().startswith('|'):
            table_lines = []
            while i < len(lines) and lines[i].strip().startswith('|'):
                table_lines.append(lines[i])
                i += 1
            process_table('\n'.join(table_lines), doc)
            continue
        
        # 处理标题
        if line.startswith('# '):
            heading = doc.add_heading(line[2:], level=1)
            heading.alignment = WD_ALIGN_PARAGRAPH.CENTER
            for run in heading.runs:
                run.font.name = '黑体'
        elif line.startswith('## '):
            heading = doc.add_heading(line[3:], level=2)
            for run in heading.runs:
                run.font.name = '黑体'
        elif line.startswith('### '):
            heading = doc.add_heading(line[4:], level=3)
            for run in heading.runs:
                run.font.name = '黑体'
        elif line.startswith('#### '):
            heading = doc.add_heading(line[5:], level=4)
            for run in heading.runs:
                run.font.name = '黑体'
        
        # 处理列表
        elif line.startswith('- '):
            p = doc.add_paragraph(style='List Bullet')
            run = p.add_run(line[2:])
            run.font.name = '宋体'
        elif line.startswith('1. '):
            p = doc.add_paragraph(style='List Number')
            run = p.add_run(line[3:])
            run.font.name = '宋体'
        
        # 处理分隔线
        elif line.strip() == '---':
            p = doc.add_paragraph()
            p.add_run('_' * 50)
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # 处理普通段落
        elif line.strip():
            p = doc.add_paragraph()
            # 设置段落格式
            p.paragraph_format.line_spacing = 1.5
            p.paragraph_format.space_after = Pt(10)
            
            # 处理加粗文本
            parts = re.split(r'(\*\*.*?\*\*)', line)
            for part in parts:
                if part.startswith('**') and part.endswith('**'):
                    run = p.add_run(part[2:-2])
                    run.bold = True
                    run.font.name = '宋体'
                else:
                    run = p.add_run(part)
                    run.font.name = '宋体'
        
        i += 1
    
    # 保存 Word 文档
    output_file = md_file.parent / "完整投标文件.docx"
    doc.save(output_file)
    print(f"已生成 Word 文档：{output_file}")
    
    # 清理配置文件
    if os.path.exists('config.json'):
        os.unlink('config.json')

if __name__ == "__main__":
    convert_md_to_word() 